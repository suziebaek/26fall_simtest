"""
doc_parser.py
워드 문서(챕터/셀 참고문서, 문제모음집)를 파싱하여 구조화된 dict 리스트로 변환.

핵심 아이디어
-------------
1. 챕터 참고문서(FG_H_챕터.docx 형태)
   "CHAPTER 0N  제목" / "CELL n   제목" 패턴을 순서대로 읽어
   {chapter_num: {"title": ..., "cells": {1: title1, 2: title2, 3: title3}}} 형태로 반환.

2. 문제모음집(FG_H_Sim_Test.docx 형태)
   - "Chapter" 단독 문단 다음에 "01  동사의 종류" 같은 챕터 타이틀이 옴 -> 챕터 경계
   - "01 다음 빈칸에...", "12", "[08-09] ..." 같은 문항 번호로 문항 경계를 잡음
   - 빨간색(EE0000/FF0000 계열) 텍스트 = 정답 표시
       - 동그라미 숫자(①~⑤) 단독이면 객관식 정답 번호
       - 문장/구(자유 텍스트)이면 서술형 정답
   - 문서 안에 섞여 있는 표(Table)는 등장 순서 그대로 문항에 붙임(참고용 원문 보존, 이미지 문항일 경우 image_promt 후보로 사용)
"""

import re
from dataclasses import dataclass, field
from docx import Document
from docx.oxml.ns import qn

CIRCLED = "①②③④⑤⑥⑦⑧⑨⑩"
RED_HEXES = {"EE0000", "FF0000", "CC0000", "C00000", "FF0033"}


def _is_red(rgb):
    if rgb is None:
        return False
    s = str(rgb).upper()
    if s in RED_HEXES:
        return True
    # 일반화: R값이 매우 높고 G,B가 낮으면 붉은색 계열로 간주
    try:
        r, g, b = int(s[0:2], 16), int(s[2:4], 16), int(s[4:6], 16)
        return r >= 180 and g <= 80 and b <= 80
    except Exception:
        return False


def _para_red_spans(paragraph):
    """문단 내에서 빨간색으로 표시된 run 텍스트들을 순서대로 반환"""
    spans = []
    for r in paragraph.runs:
        color = r.font.color.rgb if (r.font.color and r.font.color.rgb) else None
        if _is_red(color) and r.text.strip():
            spans.append(r.text)
    return spans


def _run_html(run):
    """run 하나를 텍스트로 변환하되, 밑줄 서식은 <u></u>로, 줄바꿈(w:br)은 <br/>로 표시."""
    text = run.text
    if not text:
        return ""
    text = text.replace("\t", " ").replace("\n", "<br/>")
    if run.font.underline:
        return f"<u>{text}</u>"
    return text


def _para_html(paragraph):
    """문단 하나를 밑줄(<u>)/줄바꿈(<br/>) 표시를 보존한 문자열로 변환."""
    return "".join(_run_html(r) for r in paragraph.runs)


_KNOWN_TAG_RE = re.compile(r"</?u>|<br/>")
CONDITION_MARKERS = ("<조건>", "〈조건〉", "＜조건＞")


def _is_condition_block(html_text):
    """표/문단 내용이 '<조건>'으로 시작하는지 (밑줄/줄바꿈 태그는 무시하고) 판단."""
    plain = _KNOWN_TAG_RE.sub("", html_text).strip()
    return plain.startswith(CONDITION_MARKERS)


def _table_to_html(table):
    """표 내용을 밑줄(<u>)/줄바꿈(<br/>) 서식을 보존해 문자열로 변환.
    병합된 셀은 python-docx에서 같은 셀 객체가 여러 번 나오므로 중복 제거."""
    seen_tc = set()
    para_htmls = []
    for row in table.rows:
        for cell in row.cells:
            tc_key = id(cell._tc)
            if tc_key in seen_tc:
                continue
            seen_tc.add(tc_key)
            for p in cell.paragraphs:
                html = _para_html(p).strip()
                if html:
                    para_htmls.append(html)
    return "<br/>".join(para_htmls)
CELL_RE = re.compile(r"^CELL\s*(\d+)\s+(.*)$", re.IGNORECASE)


def _parse_chapter_paragraphs(doc):
    """FG_H_챕터.docx 같은, 'CHAPTER 0N 제목' / 'CELL n 제목' 문단 패턴 형식."""
    result = {}
    cur_chapter = None
    for p in doc.paragraphs:
        t = p.text.strip()
        if not t:
            continue
        m_chapter = re.match(r"^CHAPTER\s+(\d{2})\s+(.*)$", t, re.IGNORECASE)
        if m_chapter:
            cur_chapter = int(m_chapter.group(1))
            result[cur_chapter] = {"title": m_chapter.group(2).strip(), "cells": {}}
            continue
        m_cell = CELL_RE.match(t)
        if m_cell and cur_chapter is not None:
            cell_num = int(m_cell.group(1))
            result[cur_chapter]["cells"][cell_num] = m_cell.group(2).strip()
            continue
    return result


def _find_table_col(header_cells, keywords):
    for i, h in enumerate(header_cells):
        h_low = h.lower()
        if any(kw.lower() in h_low for kw in keywords):
            return i
    return None


def _parse_chapter_table(table):
    """E_챕터.docx 같은, 'CH | 목차 | Cell' 표 형식.
    한 챕터가 여러 행에 걸쳐 반복되고(CH 값이 같은 행이 여러 개), 각 행이 하나의 CELL을 나타냄.
    CELL 번호는 문서에 없으므로 같은 챕터 안에서 등장 순서대로 1,2,3...을 부여."""
    rows = [[c.text.strip() for c in r.cells] for r in table.rows]
    if not rows:
        return {}
    header = rows[0]
    ch_idx = _find_table_col(header, ["ch", "chapter", "챕터"])
    title_idx = _find_table_col(header, ["목차", "title", "제목"])
    cell_idx = _find_table_col(header, ["cell", "셀"])
    if ch_idx is None or cell_idx is None:
        return {}

    result = {}
    cell_counters = {}
    for row in rows[1:]:
        if len(row) <= max(ch_idx, cell_idx):
            continue
        ch_raw = row[ch_idx].strip()
        cell_raw = row[cell_idx].strip()
        if not ch_raw or not cell_raw:
            continue
        m = re.search(r"(\d+)", ch_raw)
        if not m:
            continue
        ch_num = int(m.group(1))
        title = row[title_idx].strip() if title_idx is not None and len(row) > title_idx else ""

        if ch_num not in result:
            result[ch_num] = {"title": title, "cells": {}}
        elif title and not result[ch_num]["title"]:
            result[ch_num]["title"] = title

        cell_counters[ch_num] = cell_counters.get(ch_num, 0) + 1
        result[ch_num]["cells"][cell_counters[ch_num]] = cell_raw
    return result


def parse_chapter_reference(path):
    """'챕터/셀 목차' 문서를 파싱. 두 가지 형식을 자동 인식:
      1) 문단 형식 (예: FG_H_챕터.docx) - 'CHAPTER 0N 제목' / 'CELL n 제목'
      2) 표 형식 (예: E_챕터_ref.docx) - 'CH | 목차 | Cell' 컬럼을 가진 표
    반환: {chapter_num(int): {"title": str, "cells": {1: str, 2: str, 3: str}}}
    """
    doc = Document(path)

    result = _parse_chapter_paragraphs(doc)
    if result:
        return result

    # 문단 형식에서 못 찾았으면 표 형식으로 재시도
    for table in doc.tables:
        t_result = _parse_chapter_table(table)
        for ch, info in t_result.items():
            entry = result.setdefault(ch, {"title": "", "cells": {}})
            if info["title"] and not entry["title"]:
                entry["title"] = info["title"]
            entry["cells"].update(info["cells"])
    return result



def _looks_like_options_line(html_text):
    """①~⑤가 있는 줄이 '진짜 보기 목록'인지, 아니면 밑줄 오류 지문처럼 문장 중간에
    동그라미 숫자가 박혀있는 것인지 구분. 조각이 5개 초과이거나 한 조각이 지나치게
    길면(문장형) 보기 목록이 아니라 지문으로 간주. (<u>/<br/> 태그가 섞여 있어도 동작)"""
    parts = [p.strip() for p in re.split(r"(?=[①②③④⑤⑥⑦⑧⑨⑩])", html_text) if p.strip()]
    if not parts or len(parts) > 5:
        return False
    plain_parts = [_KNOWN_TAG_RE.sub("", p) for p in parts]
    return all(len(p.split()) <= 12 for p in plain_parts)


def _iter_block_items(doc):
    """문서 본문(paragraph, table)을 실제 등장 순서대로 순회"""
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            from docx.text.paragraph import Paragraph
            yield "p", Paragraph(child, doc)
        elif child.tag == qn("w:tbl"):
            from docx.table import Table
            yield "tbl", Table(child, doc)


def _table_looks_like_options(html_text):
    """표 전체가 사실은 지문이 아니라 보기(①~⑤) 목록인 경우를 감지
    (지문 없이 표 안에 선지가 바로 들어있는 케이스)."""
    plain = _KNOWN_TAG_RE.sub("", html_text)
    if not any(c in plain for c in CIRCLED):
        return False
    return _looks_like_options_line(html_text)


def _split_into_options(html_text):
    parts = re.split(r"(?=[①②③④⑤⑥⑦⑧⑨⑩])", html_text)
    cleaned = []
    for part in parts:
        part = re.sub(r"(<br/>)+$", "", part).strip()
        if part:
            cleaned.append(part)
    return cleaned


QNUM_RE = re.compile(r"^(\d{1,2})(?:\s+(.*))?$")
QRANGE_RE = re.compile(r"^\[(\d{1,2})-(\d{1,2})\]\s*(.*)$")
CHAPTER_MARK_RE = re.compile(r"^(\d{2})\s+(.+)$")


def parse_question_bank(path, default_chapter_num=1, default_chapter_title=""):
    """
    문제모음집 워드 문서를 파싱해서 챕터별 문항 리스트를 반환.
    default_chapter_num/default_chapter_title: 문서 안에 'Chapter' 헤더가 전혀
    없는 단일 챕터 문서(예: E레벨처럼 챕터 하나만 다루는 문제집)를 위한 기본값.
    문서 안에 'Chapter' 헤더가 실제로 나오면 그 값으로 자동 갱신됨(H레벨처럼
    여러 챕터를 한 파일에 담은 경우).
    반환: list of dict, 각 dict:
      {
        "chapter_num": int,
        "chapter_title": str,
        "q_num": int,
        "group_range": (start,end) or None,
        "shared_instruction": str or "",   # [08-09] 같은 그룹 공통 지시문
        "instruction": str,                # 문항 지시문(질문) - 밑줄(<u>) 보존
        "raw_lines": [str, ...],           # 지문/보기 외 원문 라인(순서 보존, <u>/<br/> 보존)
        "passage_tables": [html_str, ...], # 표로 된 지문(<조건>이 아닌 것) - text_passage로 감
        "prompt_tables": [html_str, ...],  # <조건>으로 시작하는 표 - text_prompt로 감
        "options": [str, ...],             # ①~⑤ 보기 (있는 경우, <u> 보존)
        "answer_marks": [str, ...],        # 빨간색으로 표시된 원문(정답 후보)
      }
    같은 그룹([N-M])에 속한 문항들은 그룹 헤더와 첫 문항 사이에 등장하는 공용 지문
    (표 또는 문단)을 각자 복사해서 가진다 (예: [11-12] 공용 지문 -> 11번, 12번 모두 포함).
    파싱은 휴리스틱이므로 반드시 Streamlit UI에서 사람이 검수해야 함.
    """
    doc = Document(path)
    items = list(_iter_block_items(doc))

    questions = []
    cur_chapter_num = default_chapter_num
    cur_chapter_title = default_chapter_title
    expecting_chapter_title = False

    cur_q = None
    cur_group_range = None
    cur_shared_instruction = ""

    # 그룹 헤더 이후 ~ 첫 문항 이전에 등장하는 공용 지문(그룹 내 모든 문항에 복사됨)
    gathering_shared = False
    shared_lines = []
    shared_passage_tables = []
    shared_prompt_tables = []

    def flush():
        if cur_q is not None:
            questions.append(cur_q)

    i = 0
    n = len(items)
    while i < n:
        kind, obj = items[i]

        if kind == "tbl":
            html = _table_to_html(obj)
            if html:
                if cur_q is not None and not cur_q["options"] and _table_looks_like_options(html):
                    # 지문 없이 표 안에 선지(①~⑤)가 바로 들어있는 경우
                    cur_q["options"].extend(_split_into_options(html))
                else:
                    is_cond = _is_condition_block(html)
                    if cur_q is not None:
                        (cur_q["prompt_tables"] if is_cond else cur_q["passage_tables"]).append(html)
                    elif gathering_shared:
                        (shared_prompt_tables if is_cond else shared_passage_tables).append(html)
                    # 그 외(문항/그룹 문맥 밖의 표: 목차 등)는 무시
            i += 1
            continue

        text = _para_html(obj).strip()
        if not text:
            i += 1
            continue

        # 챕터 헤더: 'Chapter' 단독 문단 다음에 '01  제목'
        if _KNOWN_TAG_RE.sub("", text).lower() == "chapter":
            expecting_chapter_title = True
            i += 1
            continue
        if expecting_chapter_title:
            m = CHAPTER_MARK_RE.match(text)
            if m:
                flush()
                cur_q = None
                cur_chapter_num = int(m.group(1))
                cur_chapter_title = _KNOWN_TAG_RE.sub("", m.group(2)).strip()
                expecting_chapter_title = False
                cur_shared_instruction = ""
                gathering_shared = False
                shared_lines, shared_passage_tables, shared_prompt_tables = [], [], []
                i += 1
                continue
            expecting_chapter_title = False  # 예상과 다르면 무시하고 계속 진행

        # 그룹 헤더: [08-09] 다음 중 ...
        m_range = QRANGE_RE.match(text)
        if m_range:
            flush()
            cur_q = None
            cur_group_range = (int(m_range.group(1)), int(m_range.group(2)))
            cur_shared_instruction = m_range.group(3).strip()
            gathering_shared = True
            shared_lines, shared_passage_tables, shared_prompt_tables = [], [], []
            i += 1
            continue

        # 문항 번호: '01 다음 ~' 또는 단독 '12'
        m_qnum = QNUM_RE.match(text)
        if m_qnum and cur_chapter_num is not None:
            qnum = int(m_qnum.group(1))
            rest = (m_qnum.group(2) or "").strip()
            # 그룹 범위 안에 있는 번호인지 확인
            in_group = cur_group_range and cur_group_range[0] <= qnum <= cur_group_range[1]
            if in_group or rest or qnum <= 30:
                # 새 문항 시작으로 판단 (오탐 방지를 위해 30 이하 숫자만 문항번호로 인정)
                flush()
                cur_q = {
                    "chapter_num": cur_chapter_num,
                    "chapter_title": cur_chapter_title,
                    "q_num": qnum,
                    "group_range": cur_group_range if in_group else None,
                    "shared_instruction": cur_shared_instruction if in_group else "",
                    "instruction": rest,
                    "raw_lines": list(shared_lines) if in_group else [],
                    "passage_tables": list(shared_passage_tables) if in_group else [],
                    "prompt_tables": list(shared_prompt_tables) if in_group else [],
                    "options": [],
                    "answer_marks": [],
                    "has_inline_markers": False,
                }
                gathering_shared = False  # 이후 내용은 이 문항 전용(공용 지문 수집은 종료)
                i += 1
                continue

        # 그 외 줄
        if cur_q is not None:
            reds = _para_red_spans(obj)
            if reds:
                cur_q["answer_marks"].extend(reds)
            # 보기(①~⑤)가 한 줄에 여러 개 있을 수 있으므로 분리 (단, 밑줄 오류 지문처럼
            # 문장 중간에 동그라미 숫자가 박힌 경우는 보기 목록이 아니라 지문으로 취급)
            if any(c in text for c in CIRCLED):
                if _looks_like_options_line(text):
                    for part in _split_into_options(text):
                        cur_q["options"].append(part)
                else:
                    cur_q["raw_lines"].append(text)
                    cur_q["has_inline_markers"] = True
            elif cur_q["options"]:
                # 이미 선지(①~⑤) 목록이 시작된 뒤에 나오는, 동그라미 번호가 없는 줄은
                # 새로운 지문이 아니라 직전 선지의 연속(예: '→ 변환된 문장')이다.
                # 다음 선지 번호가 나오기 전까지는 모두 같은 선지로 합친다.
                cur_q["options"][-1] = cur_q["options"][-1] + "<br/>" + text
            else:
                cur_q["raw_lines"].append(text)
        elif gathering_shared:
            # 그룹 헤더 이후, 첫 문항 전에 나오는 공용 지문 문단(표가 아닌 경우)
            shared_lines.append(text)
        i += 1

    flush()
    return questions
